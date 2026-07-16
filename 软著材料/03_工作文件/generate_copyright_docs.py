from __future__ import annotations

import csv
import os
import shutil
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
OUT = ROOT / "软著材料"
SCREEN_DIR = OUT / "01_界面截图"
DOC_DIR = OUT / "02_提交文档"
SOURCE_SCREEN_DIR = ROOT / "docs" / "prototype-v1-raster"
SOFTWARE_NAME = "赛小蜂篮球赛事管理软件"
VERSION = "V1.0"
FULL_NAME = f"{SOFTWARE_NAME}{VERSION}"


ROUTES = [
    ("pages/login/index", "登录页"),
    ("pages/home/index", "工作台"),
    ("pages/quick-match/index", "快捷比赛设置"),
    ("pages/scorer/index", "比赛设置/竖屏计分"),
    ("pages/scorer-board/index", "PAD横屏比赛控制台"),
    ("pages/products/index", "产品与权益"),
    ("pages/mc-system/index", "AI播音员与音效"),
    ("pages/data/index", "数据中心"),
    ("pages/education/index", "教务中心"),
    ("pages/stats-scorer/index", "技术统计计分"),
    ("pages/order/index", "订单确认"),
    ("pages/pay-result/index", "支付结果"),
    ("pages/tournament/index", "赛事列表"),
    ("pages/tournament-create/index", "创建赛事"),
    ("pages/tournament-detail/index", "赛事详情与场次"),
    ("pages/game-detail/index", "比赛详情"),
    ("pages/team/index", "球员库/球队管理"),
    ("pages/team-create/index", "创建球队"),
    ("pages/player-add/index", "创建球员"),
    ("pages/mc-settings/index", "MC音效设置"),
    ("pages/mine/index", "我的"),
]


SCREEN_SOURCES = [
    ("01-登录页.png", "01_登录页.png", "登录与游客体验"),
    ("18-工作台六栏菜单简化版.png", "02_工作台.png", "机构端工作台"),
    ("03-快速比赛设置.png", "03_快捷比赛设置.png", "快捷比赛设置"),
    ("saixiaofeng-scoreboard-final-imagegen-v4-bench-extended.png", "04_PAD横屏比赛控制台.png", "PAD横屏计分板"),
    ("05-MC音效抽屉.png", "05_MC音效控制.png", "计分板MC音效控制"),
    ("06-我的页面.png", "06_我的.png", "我的与账号资料"),
    ("07-MC音效设置.png", "07_MC音效设置.png", "MC音效设置"),
    ("08-PC管理后台.png", "08_PC管理后台.png", "PC管理后台"),
    ("19-赛事页面六栏菜单版.png", "09_赛事列表.png", "赛事列表"),
    ("10-赛事详情场次.png", "10_赛事详情与场次.png", "赛事详情与场次"),
    ("20-球员页面六栏菜单版.png", "11_球员库.png", "球员库"),
    ("21-教务中心六栏菜单版.png", "12_教务中心.png", "教务中心"),
    ("22-数据中心六栏菜单版.png", "13_数据中心.png", "数据中心"),
]


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_run_font(run, name="微软雅黑", size=10.5, bold=False, color=None) -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def setup_doc(doc: Document, landscape=False) -> None:
    sec = doc.sections[0]
    sec.top_margin = Cm(2.2)
    sec.bottom_margin = Cm(2.0)
    sec.left_margin = Cm(2.2)
    sec.right_margin = Cm(2.0)
    if landscape:
        sec.orientation = 1
        sec.page_width, sec.page_height = sec.page_height, sec.page_width
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "微软雅黑"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.line_spacing = 1.35
    normal.paragraph_format.space_after = Pt(4)
    for style_name, size in [("Title", 26), ("Heading 1", 16), ("Heading 2", 13), ("Heading 3", 11)]:
        style = styles[style_name]
        style.font.name = "微软雅黑"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor(230, 92, 0) if style_name != "Title" else RGBColor(35, 35, 35)


def add_field(paragraph, field_code: str) -> None:
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = field_code
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char1, instr_text, fld_char2])


def add_header_footer(doc: Document, title=FULL_NAME) -> None:
    for sec in doc.sections:
        p = sec.header.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(f"{title}  ·  软件文档")
        set_run_font(r, size=9, color=(95, 95, 95))
        f = sec.footer.paragraphs[0]
        f.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = f.add_run("第 ")
        set_run_font(r, size=9)
        add_field(f, "PAGE")
        r = f.add_run(" 页")
        set_run_font(r, size=9)


def add_cover(doc: Document, subtitle: str) -> None:
    for _ in range(4):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(SOFTWARE_NAME)
    set_run_font(r, size=26, bold=True, color=(230, 92, 0))
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(VERSION)
    set_run_font(r, size=18, bold=True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(36)
    r = p.add_run(subtitle)
    set_run_font(r, size=22, bold=True)
    for _ in range(7):
        doc.add_paragraph()
    for label, value in [("软件名称", SOFTWARE_NAME), ("版本号", VERSION), ("著作权人", "[待填写]"), ("文档日期", "2026年7月")]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(f"{label}：{value}")
        set_run_font(r, size=12)
    doc.add_page_break()


def add_caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    set_run_font(r, size=9, color=(95, 95, 95))


def copy_screens() -> list[tuple[Path, str]]:
    SCREEN_DIR.mkdir(parents=True, exist_ok=True)
    result = []
    for src_name, dst_name, caption in SCREEN_SOURCES:
        src = SOURCE_SCREEN_DIR / src_name
        dst = SCREEN_DIR / dst_name
        # 用户可直接在 01_界面截图 中替换最终运行页面。
        # 重新生成文档时必须优先使用这些最终图，不得被旧原型覆盖。
        if not dst.exists():
            shutil.copy2(src, dst)
        result.append((dst, caption))
    # 附加保留项目中的全部 1.0 原型图，便于申报人选图。
    all_dir = SCREEN_DIR / "附录_项目现有全部界面图"
    all_dir.mkdir(exist_ok=True)
    for src in sorted(SOURCE_SCREEN_DIR.glob("*.png")):
        shutil.copy2(src, all_dir / src.name)
    return result


def build_screenshot_manifest(screens: list[tuple[Path, str]]) -> None:
    md = [f"# {FULL_NAME}界面截图清单", "", "## 一、已整理提交用截图", ""]
    for idx, (path, caption) in enumerate(screens, 1):
        md.append(f"{idx}. {caption}：`{path.name}`")
    md += ["", "## 二、小程序已注册页面（21个）", ""]
    for idx, (route, title) in enumerate(ROUTES, 1):
        md.append(f"{idx}. {title}：`/{route}`")
    md += [
        "",
        "## 三、说明",
        "",
        "- 提交用图优先选用项目内已验证的1.0界面图集。",
        "- `附录_项目现有全部界面图`保留了项目当前所有原型/实现图，可供最终申报选图。",
        "- 微信开发者工具的运行时自动截图需先在工具中开启服务端口；开启后可按上述路由批量替换。",
    ]
    (SCREEN_DIR / "界面截图清单.md").write_text("\n".join(md), encoding="utf-8")
    with (SCREEN_DIR / "已注册页面路由清单.csv").open("w", newline="", encoding="utf-8-sig") as f:
        writer = csv.writer(f)
        writer.writerow(["序号", "页面名称", "路由"])
        for idx, (route, title) in enumerate(ROUTES, 1):
            writer.writerow([idx, title, "/" + route])


def build_info_form(total_lines: int) -> Path:
    doc = Document()
    setup_doc(doc)
    add_header_footer(doc, FULL_NAME)
    add_cover(doc, "软件信息采集表（申报草案）")
    doc.add_heading("填写说明", level=1)
    p = doc.add_paragraph("本表已根据当前项目代码、功能与运行形态预填。涉及申请主体、日期和证件的字段无法从代码确认，已以“[待填写]”标记，提交前必须核实。")
    p.runs[0].font.color.rgb = RGBColor(180, 30, 30)

    rows = [
        ("软件名称", SOFTWARE_NAME),
        ("版本号", VERSION),
        ("软件简称", "赛小蜂篮球"),
        ("软件分类", "应用软件 / 小程序"),
        ("软件作品说明", "原创"),
        ("公司成立时间", "[待填写]"),
        ("开发完成日期", "[待填写]"),
        ("发表状态", "[已发表/未发表，待确认]"),
        ("开发方式", "[独立开发，待确认]"),
        ("著作权人", "[待填写公司或个人全称]"),
        ("证件类型及号码", "[待填写]"),
        ("国籍/省市", "中国 / [待填写]"),
        ("权利取得方式", "原始取得"),
        ("权利范围", "全部权利"),
        ("交存方式", "一般交存：源程序前连续30页和后连续30页；一种文档"),
        ("开发硬件环境", "x64架构开发电脑，8GB以上内存，具备网络环境"),
        ("运行硬件环境", "支持微信的智能手机或平板电脑，计分控制台推荐PAD横屏"),
        ("开发操作系统", "Windows 11"),
        ("开发环境/工具", "微信开发者工具、Node.js、Visual Studio Code、Git"),
        ("运行平台/操作系统", "微信小程序平台；iOS、Android、HarmonyOS等微信支持环境"),
        ("运行支撑环境", "微信客户端及小程序基础库；云开发/云函数服务"),
        ("编程语言", "JavaScript、WXML、WXSS、JSON"),
        ("源程序量", f"约 {total_lines:,} 行（以本次统计为准）"),
        ("开发目的", "为篮球赛事组织者和教培机构提供便捷、专业的赛事创建、现场计分和资料管理工具。"),
        ("面向领域/行业", "体育赛事、篮球教育培训、场馆和青少年篮球机构"),
        ("软件的主要功能", "系统支持微信登录、快捷比赛、赛事创建与场次管理、球队与球员库、PAD横屏现场计分、比分及犯规暂停记录、操作撤销、阵容选择、MC音效和比赛数据管理，并为教务与数据中心后续扩展提供统一入口。"),
        ("技术特点", "采用微信原生小程序架构和云函数，面向PAD横屏场景优化计分控制台；以事件日志支持上一步撤销，通过六栏导航统一赛事、球员、教务和数据入口。"),
        ("软件选择", "小程序、云计算软件、大数据软件（最终选择请申报人确认）"),
    ]
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.columns[0].width = Cm(5)
    table.columns[1].width = Cm(11.5)
    hdr = table.rows[0].cells
    hdr[0].text, hdr[1].text = "字段", "预填内容"
    set_repeat_table_header(table.rows[0])
    for c in hdr:
        set_cell_shading(c, "F4B183")
        for r in c.paragraphs[0].runs:
            set_run_font(r, bold=True)
    for label, value in rows:
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = value
        cells[0].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        if "待" in value:
            for r in cells[1].paragraphs[0].runs:
                r.font.color.rgb = RGBColor(180, 30, 30)
    path = DOC_DIR / f"{FULL_NAME}_软件信息采集表.docx"
    doc.save(path)
    return path


def build_features() -> Path:
    doc = Document()
    setup_doc(doc)
    add_header_footer(doc)
    add_cover(doc, "软件功能和技术特点说明")
    sections = [
        ("一、软件概述", [
            f"{FULL_NAME}是面向篮球赛事组织者、教培机构、场馆和计分人员的微信小程序。软件以“赛事MVP+现场计分”为1.0核心，将比赛创建、场次组织、球队球员资料、现场操作和MC音效纳入统一工作台。",
            "系统兼顾临时比赛和正式赛事：快捷比赛可在不绑定双方球员资料的情况下直接开赛；专业赛事可使用平台内球队、球员库和场次信息进行精细管理。",
        ]),
        ("二、主要功能", [
            "1. 登录与身份入口：支持微信授权登录和游客体验，向机构与现场工作人员提供统一入口。",
            "2. 工作台：提供创建赛事、快捷比赛、创建球队和创建球员等高频功能。",
            "3. 快捷比赛：设置主客队、每节时长、比赛节数和计时方式，适合试听课、内部对抗和临时单场比赛。",
            "4. 赛事管理：支持赛事创建、状态筛选、赛事详情、场次设置、主客队对阵与计分入口。",
            "5. 球队与球员库：建立球队和球员基础资料，支持搜索、筛选和赛事阵容选择。",
            "6. PAD横屏计分控制台：集成比分、节次、比赛时间、24秒、球权、犯规、暂停、换人、减1分和上一步撤销。",
            "7. 阵容管理：主客队均可从球员库搜索添加首发和替补，阵容区域支持独立上下滚动与删除。",
            "8. MC音效：提供2分、3分、罚球进/失等常用音效键以及6个自定义音效键。",
            "9. 操作日志：主客队关键操作实时记录，便于现场核对和上一步撤销。",
            "10. 教务与数据中心：1.0作为后续版本的统一展示入口，保留多校区、排课、课消、评价和赛事数据等扩展能力。",
        ]),
        ("三、技术特点", [
            "1. 原生小程序实现：界面与交互使用WXML、WXSS和JavaScript实现，具有体积较小、启动快和微信生态兼容性好的特点。",
            "2. 云端业务支撑：使用云函数承载登录、用户资料、赛果、订单、支付、音效库和权益管理等业务。",
            "3. 横竖屏分场景设计：普通管理页面面向手机竖屏，现场计分控制台面向PAD横屏，在同一产品内适配不同操作环境。",
            "4. 可撤销计分操作：将比分、犯规、暂停等现场操作纳入历史记录，为“撤销上一步”和现场复核提供基础。",
            "5. 球员与球队解耦：球员库与球队可分别管理，只在具体比赛阵容中按需添加，兼容临时参赛和长期资料累积。",
            "6. 统一六栏导航：工作台、赛事、球员、教务、数据和我的形成统一信息架构，便于1.0与2.0能力平滑扩展。",
        ]),
        ("四、运行环境与性能", [
            "软件运行于支持微信小程序的智能手机或平板电脑。日常管理界面建议使用手机竖屏，现场计分建议使用较大屏幕的PAD横屏。",
            "系统关键计分操作在本地界面即时反馈；需要云端数据的登录、赛事、赛果、音效库等功能依赖网络环境。",
        ]),
        ("五、安全与可维护性", [
            "敏感配置与业务代码分离，支付密钥、管理员令牌等内容通过环境变量维护，不写入小程序前端。",
            "项目按页面、公共组件、静态资源、工具函数和云函数进行模块化组织，支持后续按赛事、教务、数据和会员权益分阶段扩展。",
        ]),
    ]
    for heading, paras in sections:
        doc.add_heading(heading, level=1)
        for text in paras:
            doc.add_paragraph(text)
    path = DOC_DIR / f"{FULL_NAME}_功能特点.docx"
    doc.save(path)
    return path


def build_manual(screens: list[tuple[Path, str]]) -> Path:
    doc = Document()
    setup_doc(doc)
    add_header_footer(doc)
    add_cover(doc, "用户手册")
    doc.add_heading("一、文档历史", level=1)
    table = doc.add_table(rows=2, cols=5)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    values = [["版本", "修改内容", "修改日期", "修改人", "审阅人"], ["V1.0", "首次编制", "2026年7月", "[待填写]", "[待填写]"]]
    for i, row in enumerate(values):
        for j, value in enumerate(row):
            table.cell(i, j).text = value
            if i == 0:
                set_cell_shading(table.cell(i, j), "F4B183")
    doc.add_heading("二、引言", level=1)
    doc.add_heading("2.1 编写目的", level=2)
    doc.add_paragraph(f"本手册用于说明{FULL_NAME}的功能、运行环境和操作流程，供赛事组织者、教培机构管理人员、计分员和系统维护人员参考。")
    doc.add_heading("2.2 适用范围", level=2)
    doc.add_paragraph("本手册适用于软件V1.0小程序端及与其配套的基础PC管理后台。教务和数据中心在1.0中主要作为功能介绍与后续版本入口。")
    doc.add_heading("2.3 运行环境", level=2)
    for line in ["客户端：安装微信的iOS、Android或HarmonyOS手机/PAD。", "网络：登录、云端赛事和音效库等功能需要网络连接。", "建议：日常管理使用竖屏，现场计分使用PAD横屏。"]:
        doc.add_paragraph(line, style="List Bullet")
    doc.add_page_break()

    chapters = [
        ("三、登录与进入系统", 0, ["打开赛小蜂篮球小程序。", "阅读并勾选用户协议与隐私政策。", "点击微信授权登录；如需快速了解功能，可使用游客体验。"]),
        ("四、工作台", 1, ["工作台展示今日比赛、待开始和已完成数量。", "快捷功能包括创建赛事、快捷比赛、创建球队和创建球员。", "底部六栏菜单可切换工作台、赛事、球员、教务、数据和我的。"]),
        ("五、创建快捷比赛", 2, ["在工作台点击“快捷比赛”或“开始设置”。", "填写或选择主队、客队，设置每节时长、比赛节数和正/倒计时。", "快捷比赛可不绑定双方球员数据，点击开始比赛即可进入计分板。"]),
        ("六、PAD横屏比赛控制台", 3, ["顶部显示赛小蜂比赛控制台、比赛名称/场次、当前节次、重置与设置。", "中部控制比赛时间、节次、球权、24秒、交换场地和结束比赛。", "左右球队区提供罚球+1、两分+2、三分+3、减1、犯规、暂停和换人。", "“撤销”用于撤销上一步可追溯操作；误点后应优先使用撤销，避免人工连续反向修改。"]),
        ("七、MC音效操作", 4, ["计分板底部提供常用音效键与自定义音效键。", "常用音效包括2分、3分、罚球进和罚球失等现场播放项。", "自定义音效区保留6个按键，可在音效设置中配置后使用。"]),
        ("八、赛事管理", 8, ["进入“赛事”查看赛事列表，按状态查看待开始、进行中或已完成赛事。", "点击创建赛事，填写赛事基础资料并保存。", "打开赛事详情可维护场次、主客队和计分入口。"]),
        ("九、赛事详情与场次", 9, ["赛事详情页集中展示赛事名称、时间、场地、参赛球队和场次。", "选择具体场次后可查看对阵信息并进入计分页。", "正式比赛可选择平台内已创建的球队，阵容再从球员库按需添加。"]),
        ("十、球员库与阵容", 10, ["进入“球员”查看球员库，支持姓名搜索和球队筛选。", "新建球员时填写姓名、号码、位置、身高、年龄等基础资料。", "计分板阵容区主客队分别维护，首发最多5人、替补最多6人；列表可独立上下滚动，删除按键用于移出本场阵容。"]),
        ("十一、教务中心", 11, ["1.0版本展示多校区、课程排课、学员签到、课消统计、课后评价和教练工资等后续能力。", "本页在1.0主要用于功能说明，不作为已开通的完整校务流程。"]),
        ("十二、数据中心", 12, ["1.0版本展示赛事数据、球员技术统计、球队对比、成长趋势、家长同步和数据海报等后续能力。", "本页在1.0主要用于功能说明；赛后个人报告不属于当前1.0必备提交能力。"]),
        ("十三、我的与MC设置", 5, ["“我的”页用于查看和修改头像、昵称和基础账号资料。", "从我的页面进入MC音效设置，可查看和配置计分板播放的音效。", "如使用自定义音效或个人声音语音包，应根据实际权益与服务状态操作。"]),
        ("十四、PC管理后台", 7, ["PC后台用于维护会员码、MC音乐库和账号权益等管理能力。", "后台属于管理人员使用界面，应根据账号权限访问，普通小程序用户无需使用。"]),
    ]
    for title, image_idx, steps in chapters:
        doc.add_heading(title, level=1)
        for step in steps:
            doc.add_paragraph(step, style="List Number")
        if 0 <= image_idx < len(screens):
            image_path, caption = screens[image_idx]
            doc.add_paragraph()
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            is_landscape = image_path.name.startswith("04_") or image_path.name.startswith("05_") or image_path.name.startswith("08_")
            p.add_run().add_picture(str(image_path), width=Cm(16.2 if is_landscape else 8.6))
            add_caption(doc, f"图：{caption}")
            if image_idx == 5 and len(screens) > 6:
                image_path, caption = screens[6]
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.add_run().add_picture(str(image_path), width=Cm(8.6))
                add_caption(doc, f"图：{caption}")
        doc.add_page_break()
    doc.add_heading("十五、常见问题", level=1)
    faq = [
        ("误操作加分后如何处理？", "优先点击“撤销”撤销上一步操作；若已进行多步其他操作，按现场记录核对后使用减1或对应控件修正。"),
        ("快捷比赛必须添加球员吗？", "不必须。快捷比赛可锁定球员区直接计分；只有需要统计具体球员数据时才需要从球员库添加。"),
        ("为什么音效无法播放？", "请检查设备媒体音量、微信音频权限、网络状态以及对应音效是否已配置。"),
        ("计分板显示不完整怎么办？", "建议使用PAD横屏，关闭系统分屏和过大的显示缩放，并确保微信已更新至正常版本。"),
    ]
    for q, a in faq:
        p = doc.add_paragraph()
        r = p.add_run("Q：" + q)
        set_run_font(r, bold=True, color=(230, 92, 0))
        doc.add_paragraph("A：" + a)
    path = DOC_DIR / f"{FULL_NAME}_用户手册.docx"
    doc.save(path)
    return path


def collect_source_lines() -> tuple[list[str], int, list[Path]]:
    roots = [ROOT / "native-dist", ROOT / "cloudfunctions", ROOT / "admin"]
    exts = {".js", ".wxml", ".wxss", ".json", ".html", ".css"}
    excluded_parts = {"node_modules", "assets", "local-assets", "dist", ".git"}
    excluded_names = {"project.private.config.json", "package-lock.json"}
    files: list[Path] = []
    total = 0
    blocks: list[list[str]] = []
    for base in roots:
        if not base.exists():
            continue
        for path in sorted(base.rglob("*")):
            if not path.is_file() or path.suffix.lower() not in exts:
                continue
            if path.name in excluded_names or any(part in excluded_parts for part in path.parts):
                continue
            try:
                text = path.read_text(encoding="utf-8")
            except UnicodeDecodeError:
                continue
            raw_lines = text.splitlines()
            if not raw_lines:
                continue
            rel = path.relative_to(ROOT).as_posix()
            block = [f"/* ===== 文件：{rel} ===== */"] + raw_lines
            blocks.append(block)
            files.append(path)
            total += len(raw_lines)
    all_lines = [line for block in blocks for line in block]
    return all_lines, total, files


def build_source_doc(all_lines: list[str]) -> Path:
    needed = 3000
    if len(all_lines) >= needed:
        selected = all_lines[:1500] + all_lines[-1500:]
    else:
        selected = all_lines[:]
    # 如代码不足3000行，完整交存；当前项目实际远超该数量。
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Cm(1.25)
    sec.bottom_margin = Cm(1.15)
    sec.left_margin = Cm(1.45)
    sec.right_margin = Cm(1.25)
    sec.header_distance = Cm(0.45)
    sec.footer_distance = Cm(0.45)
    header = sec.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = header.add_run(f"{SOFTWARE_NAME}  {VERSION}  源程序")
    set_run_font(r, size=8.5, bold=True)
    footer = sec.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = footer.add_run("第 ")
    set_run_font(r, size=8)
    add_field(footer, "PAGE")
    r = footer.add_run(" 页")
    set_run_font(r, size=8)
    normal = doc.styles["Normal"]
    normal.font.name = "Consolas"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(6.8)
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.line_spacing = 1.0
    for page_no in range((len(selected) + 49) // 50):
        chunk = selected[page_no * 50:(page_no + 1) * 50]
        for i, line in enumerate(chunk, page_no * 50 + 1):
            safe = line.replace("\t", "    ")
            if len(safe) > 142:
                safe = safe[:139] + "..."
            p = doc.add_paragraph()
            p.paragraph_format.keep_together = True
            r = p.add_run(f"{i:04d}  {safe}")
            set_run_font(r, name="Consolas", size=6.8)
        if page_no < (len(selected) + 49) // 50 - 1:
            doc.add_page_break()
    path = DOC_DIR / f"{FULL_NAME}_源代码.docx"
    doc.save(path)
    return path


def build_catalog(screens: list[tuple[Path, str]]) -> Path:
    doc = Document()
    setup_doc(doc)
    add_header_footer(doc)
    add_cover(doc, "软件界面图集")
    doc.add_heading("一、页面路由清单", level=1)
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for idx, value in enumerate(["序号", "页面", "路由"]):
        table.cell(0, idx).text = value
        set_cell_shading(table.cell(0, idx), "F4B183")
    for idx, (route, title) in enumerate(ROUTES, 1):
        cells = table.add_row().cells
        cells[0].text, cells[1].text, cells[2].text = str(idx), title, "/" + route
    doc.add_page_break()
    doc.add_heading("二、主要界面", level=1)
    for idx, (path, caption) in enumerate(screens, 1):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        landscape = path.name.startswith("04_") or path.name.startswith("05_") or path.name.startswith("08_")
        p.add_run().add_picture(str(path), width=Cm(16.5 if landscape else 9.0))
        add_caption(doc, f"图{idx}  {caption}")
        if idx != len(screens):
            doc.add_page_break()
    path = DOC_DIR / f"{FULL_NAME}_界面图集.docx"
    doc.save(path)
    return path


def write_readme(paths: list[Path], total_lines: int, source_files: int) -> None:
    lines = [
        f"# {FULL_NAME}软著材料说明",
        "",
        "## 已生成",
        "",
    ]
    lines.extend(f"- `{path.relative_to(OUT).as_posix()}`" for path in paths)
    lines += [
        "",
        "## 代码统计",
        "",
        f"- 纳入统计的源文件：{source_files} 个",
        f"- 实际源程序行数：约 {total_lines:,} 行",
        "- 源代码文档：一般交存前30页+后30页，每页50行，共60页。",
        "",
        "## 提交前必须补充",
        "",
        "- 著作权人法定全称、证件类型和证件号码。",
        "- 公司成立日期、开发完成日期、首次发表状态/日期/地点。",
        "- 确认开发方式、权利取得方式和权利范围。",
        "- 确认最终申报软件名称与简称未被他人占用。",
        "",
        "## 截图说明",
        "",
        "当前提交稿使用项目内已保存的1.0原型/实现图。微信开发者工具开启服务端口后，可按 `01_界面截图/已注册页面路由清单.csv` 对21个注册页面进行运行时批量截图与替换。",
    ]
    (OUT / "README_提交前必读.md").write_text("\n".join(lines), encoding="utf-8")


def main() -> None:
    SCREEN_DIR.mkdir(parents=True, exist_ok=True)
    DOC_DIR.mkdir(parents=True, exist_ok=True)
    screens = copy_screens()
    build_screenshot_manifest(screens)
    all_lines, total_lines, source_files = collect_source_lines()
    paths = [
        build_info_form(total_lines),
        build_features(),
        build_manual(screens),
        build_source_doc(all_lines),
        build_catalog(screens),
    ]
    write_readme(paths, total_lines, len(source_files))
    print(f"screens={len(screens)} total_source_lines={total_lines} source_files={len(source_files)}")
    for path in paths:
        print(path)


if __name__ == "__main__":
    main()
